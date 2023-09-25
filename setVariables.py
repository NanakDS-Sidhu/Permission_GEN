from docx import Document
from docx2pdf import convert

def replace_text_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for key, value in replacements.items():
            # print(key,value)
            if key in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)

# Replace placeholders in the document


def main():
        document = Document("template.docx")
        paras=document.paragraphs
        for i in range(len(paras)):
            paras[i]=paras[i].text
        # print(paras)
        
        replacement_values = {
            "secy_SID": "Your Secretary's ID",
            "secy_name": "Your Secretary's Name",
            "start_time": "Start Time",
            "end_time": "End Time",
            "start_date": "Start Date",
            "end_date": "End Date"
        }

        replace_text_in_paragraphs(document.paragraphs, replacement_values)
                
        document.save("temp.docx")
        convert("temp.docx", 'output.pdf')  

if __name__=="__main__":
    main()